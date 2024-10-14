# Import necessary libraries
import os
import docx
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import mac_morpho
from nltk.tag import UnigramTagger, BigramTagger
import spacy
import unicodedata

# Download necessary NLTK resources
nltk.download('punkt')
nltk.download('mac_morpho')

# Load the Portuguese model for spacy
nlp = spacy.load('pt_core_news_lg')

# Function to remove accented characters from text
def remove_accented_chars(text):
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8', 'ignore')
    return text

# Function to extract text from DOCX file
def extract_text_from_docx(docx_path):
    try:
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Erro ao ler o arquivo DOCX: {e}")

# Function to tokenize text into sentences
def tokenize_sentences(text):
    return sent_tokenize(text, language='portuguese')

# Function to tokenize sentences into words and identify verbs
def extract_verbs(sentences):
    # Train a POS tagger using the mac_morpho corpus
    train_sents = mac_morpho.tagged_sents()
    unigram_tagger = UnigramTagger(train_sents)
    bigram_tagger = BigramTagger(train_sents, backoff=unigram_tagger)
    
    verbs = []
    for sentence in sentences:
        words = word_tokenize(sentence, language='portuguese')
        tagged_words = bigram_tagger.tag(words)
        verbs.extend([word for word, tag in tagged_words if tag and tag.startswith('V')])
    return verbs

# Function to convert verbs to their infinitive form using spacy
def convert_to_infinitive(verbs):
    infinitive_verbs = []
    for verb in verbs:
        doc = nlp(verb)
        for token in doc:
            if token.pos_ == 'VERB':
                infinitive_verbs.append(token.lemma_)
    return infinitive_verbs

# Function to list DOCX files in the Books directory
def list_books(directory):
    try:
        return [f for f in os.listdir(directory) if f.endswith('.docx')]
    except FileNotFoundError:
        print(f"Pasta '{directory}' não encontrada.")
        return []

# Function to display a menu to select a DOCX file
def select_book(books):
    if not books:
        print("Nenhum livro encontrado.")
        return None
    
    print("Livros disponíveis para processar:")
    for i, book in enumerate(books, 1):
        print(f"{i}. {book}")

    while True:
        try:
            choice = int(input("\nSelecione o número do livro que deseja processar: "))
            if 1 <= choice <= len(books):
                return books[choice - 1]
            else:
                print("Escolha inválida. Tente novamente.")
        except ValueError:
            print("Entrada inválida. Insira um número.")

# Function to replace verbs in the text with their infinitive forms using spacy
def replace_verbs_with_infinitive(text):
    doc = nlp(text)
    modified_text = []
    for token in doc:
        if token.pos_ == 'VERB':
            modified_text.append(token.lemma_)
        else:
            modified_text.append(token.text)
    return ' '.join(modified_text)

# Main function to process the document, replace verbs, and export the modified text
def main():
    books_dir = os.path.join(os.path.dirname(__file__), 'Books')
    books = list_books(books_dir)
    selected_book = select_book(books)

    if selected_book:
        # Caminho completo do livro selecionado
        docx_path = os.path.join(books_dir, selected_book)

        # Extrair o texto do DOCX
        text = extract_text_from_docx(docx_path)
        
        # Normalize the text by removing accents and extra spaces
        text = remove_accented_chars(text)
        text = ' '.join(text.split()).lower()
        print(f"Normalized text: {text[:100]}...")  # Debug statement to check normalization
        
        sentences = tokenize_sentences(text)
        verbs = extract_verbs(sentences)
        infinitive_verbs = convert_to_infinitive(verbs)

        # Replace verbs with their infinitive forms
        modified_text = replace_verbs_with_infinitive(text)

        # Exportar os verbos no infinitivo como um arquivo .txt
        output_file = selected_book.replace('.docx', '_verbs.txt')
        with open(output_file, 'w', encoding='utf-8') as f:
            for verb in infinitive_verbs:
                f.write(f"{verb}\n")
        print(f"\nLista de verbos no infinitivo salva como '{output_file}'.")

        # Save the modified text back to a new DOCX file
        modified_docx_path = selected_book.replace('.docx', '_modified.docx')
        doc = docx.Document()
        for paragraph in modified_text.split('\n'):
            doc.add_paragraph(paragraph)
        doc.save(modified_docx_path)
        print(f"\nTexto modificado salvo como '{modified_docx_path}'.")

# Example usage
if __name__ == "__main__":
    main()